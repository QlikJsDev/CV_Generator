"""
Generator for the Select Advisory CV template.

Layout:
- Linear (no floating shapes)
- Word header section contains the Select Advisory logo (preserved automatically)
- Body: Heading 1 (name), Heading 2 (sections), List Paragraph (bullets),
        Job Dates / Job Details / White text (experience entries)
"""

import io
from docx import Document
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK

W  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
WP = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'


def generate_select_advisory(data, template_path):
    doc = Document(template_path)
    _rebuild_body(doc, data)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── body rebuild ──────────────────────────────────────────────────────────────

def _rebuild_body(doc, data):
    body = doc.element.body
    children = list(body)

    anchor_indices = {
        i for i, child in enumerate(children)
        if child.tag == f'{{{W}}}p' and child.findall(f'.//{{{WP}}}anchor')
    }

    to_remove = [
        child for i, child in enumerate(children)
        if child.tag in (f'{{{W}}}p', f'{{{W}}}tbl') and i not in anchor_indices
    ]
    for child in to_remove:
        body.remove(child)

    _add_content(doc, data)


def _safe_add(doc, text, style):
    try:
        return doc.add_paragraph(text, style=style)
    except Exception:
        return doc.add_paragraph(text)


# ── content builder ───────────────────────────────────────────────────────────

def _add_content(doc, data):
    first = data.get('firstName', '')
    last  = data.get('lastName', '')
    name  = f'{first} {last}'.strip()

    # Name (Heading 1)
    _safe_add(doc, name, 'Heading 1')

    # Personal info block: birthdate + roles joined with soft line breaks
    info_lines = []
    if data.get('birthDate'):
        info_lines.append(data['birthDate'])
    for title in data.get('titles', []):
        if title.strip():
            info_lines.append(title)

    if info_lines:
        p = _safe_add(doc, '', 'Normal')
        for i, line in enumerate(info_lines):
            if i > 0:
                p.add_run().add_break(WD_BREAK.LINE)
            p.add_run(line)

    # Bio
    bio = data.get('bio', '')
    if bio:
        _safe_add(doc, bio, 'Normal')

    # ── Professional summary ─────────────────────────────────────────────────
    career = data.get('careerTimeline', [])
    if career:
        _safe_add(doc, 'Professional summary', 'Heading 2')
        for entry in career:
            dates   = entry.get('dates', '')
            role    = entry.get('role', '')
            company = entry.get('company', '')
            line = f'{dates}: {role}' if dates else role
            if company:
                line += f' at {company}'
            _safe_add(doc, line, 'Normal')

    # ── Know-how (placeholder) ───────────────────────────────────────────────
    # Only added if there's content
    knowhow = data.get('knowhow', '').strip()
    if knowhow:
        _safe_add(doc, 'Know-how', 'Heading 2')
        _safe_add(doc, knowhow, 'Normal')

    # ── Technical expertise ──────────────────────────────────────────────────
    tech_skills = data.get('technicalSkills', [])
    if tech_skills:
        _safe_add(doc, 'Technical expertise', 'Heading 2')
        for skill in tech_skills:
            if skill.strip():
                _safe_add(doc, skill, 'List Paragraph')

    # ── Personal skills & areas of expertise (table) ─────────────────────────
    personal_skills = data.get('personalSkills', [])
    areas           = data.get('areasOfExpertise', [])

    if personal_skills or areas:
        try:
            table = doc.add_table(rows=1, cols=2)
            # Remove borders for a clean look
            for cell in table.rows[0].cells:
                for border_name in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV'):
                    _set_cell_border(cell, border_name, 'none')

            c0, c1 = table.cell(0, 0), table.cell(0, 1)

            _table_heading(c0, 'Personal skills')
            for s in personal_skills:
                if s.strip():
                    _safe_add_to_cell(c0, s, 'List Paragraph')

            _table_heading(c1, 'Areas of expertise')
            for a in areas:
                if a.strip():
                    _safe_add_to_cell(c1, a, 'List Paragraph')
        except Exception:
            # Fallback: flat list
            if personal_skills:
                _safe_add(doc, 'Personal skills', 'Heading 2')
                for s in personal_skills:
                    _safe_add(doc, s, 'List Paragraph')
            if areas:
                _safe_add(doc, 'Areas of expertise', 'Heading 2')
                for a in areas:
                    _safe_add(doc, a, 'List Paragraph')

    # ── Education & training ─────────────────────────────────────────────────
    education = data.get('education', [])
    if education:
        _safe_add(doc, 'Education & training', 'Heading 2')
        for edu in reversed(education):  # oldest first, like the template
            years       = edu.get('years', '')
            degree      = edu.get('degree', '')
            institution = edu.get('institution', '')
            line = f'{years}: {degree}' if years else degree
            if institution:
                line += f' — {institution}'
            _safe_add(doc, line, 'List Paragraph')

    # ── Certifications ───────────────────────────────────────────────────────
    certs = data.get('certifications', [])
    if certs:
        _safe_add(doc, 'Certifications', 'Heading 2')
        for cert in certs:
            year = cert.get('year', '')
            name = cert.get('name', '')
            label = f'{name} ({year})' if year else name
            _safe_add(doc, label, 'List Paragraph')

    # ── Languages ────────────────────────────────────────────────────────────
    languages = data.get('languages', [])
    if languages:
        _safe_add(doc, 'Languages', 'Heading 2')
        for lang in languages:
            line = f"{lang.get('language', '')}: {lang.get('proficiency', '')}"
            _safe_add(doc, line, 'List Paragraph')

    # ── Experiences ──────────────────────────────────────────────────────────
    experiences = data.get('experiences', [])
    sa_exps  = [e for e in experiences if e.get('category', 'select_advisory') != 'pre_advisory']
    pre_exps = [e for e in experiences if e.get('category', 'select_advisory') == 'pre_advisory']

    if sa_exps:
        _safe_add(doc, 'Professional experience with Select Advisory / Beyond Data', 'Heading 2')
        for exp in sa_exps:
            _add_experience(doc, exp)

    if pre_exps:
        _safe_add(doc, 'Professional experience before Select Advisory / Beyond Data', 'Heading 2')
        for exp in pre_exps:
            _add_experience(doc, exp)

    if experiences and not sa_exps and not pre_exps:
        _safe_add(doc, 'Professional experience', 'Heading 2')
        for exp in experiences:
            _add_experience(doc, exp)


def _add_experience(doc, exp):
    dates       = exp.get('dates', '')
    title       = exp.get('title', '')
    company     = exp.get('company', '')
    description = exp.get('description', [])
    tools       = exp.get('tools', '')

    if dates:
        _safe_add(doc, dates, 'Job Dates')

    job_line = title
    if company:
        job_line += f' – {company}'
    if job_line:
        _safe_add(doc, job_line, 'Job Details')

    body_style = 'White text'
    for desc in description:
        if desc.strip():
            try:
                doc.add_paragraph(desc, style=body_style)
            except Exception:
                doc.add_paragraph(desc)

    if tools:
        try:
            p = doc.add_paragraph(style=body_style)
            p.add_run(f'Tools: {tools}').bold = True
        except Exception:
            doc.add_paragraph(f'Tools: {tools}')


# ── table helpers ─────────────────────────────────────────────────────────────

def _table_heading(cell, text):
    p = cell.paragraphs[0]
    try:
        p.style = cell._element.getparent().getparent().getparent().styles['Heading 2']
    except Exception:
        pass
    p.text = text


def _safe_add_to_cell(cell, text, style):
    try:
        p = cell.add_paragraph(text, style=style)
    except Exception:
        p = cell.add_paragraph(text)
    return p


def _set_cell_border(cell, border_name, val):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(f'{{{W}}}tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    border = OxmlElement(f'w:{border_name}')
    border.set(f'{{{W}}}val', val)
    tcBorders.append(border)
