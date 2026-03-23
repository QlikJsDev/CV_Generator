"""
Generator for the Beyond Data Group CV template.

Layout:
- Floating header text box: name + 2 titles
- Floating left sidebar text box: skills + languages
- Body paragraphs (with left indent): summary, certifications, education, experiences
"""

import io
import copy
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

W   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
WP  = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
WPS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'

# ── helpers ──────────────────────────────────────────────────────────────────

def _all_text(el):
    return ''.join(t.text or '' for t in el.findall(f'.//{{{W}}}t'))


def _find_shapes(doc):
    """Return list of (anchor, txbxContent) for every text box in the document."""
    result = []
    for anchor in doc.element.body.findall(f'.//{{{WP}}}anchor'):
        txbx = anchor.find(f'.//{{{WPS}}}txbx')
        if txbx is not None:
            tc = txbx.find(f'{{{W}}}txbxContent')
            if tc is not None:
                result.append((anchor, tc))
    return result


def _replace_para_text(para, new_text):
    """Replace all runs in a paragraph with a single run carrying the first run's rPr."""
    runs = para.findall(f'{{{W}}}r')
    rPr = None
    if runs:
        rPr = runs[0].find(f'{{{W}}}rPr')
        for r in runs:
            para.remove(r)

    r = OxmlElement('w:r')
    if rPr is not None:
        r.append(copy.deepcopy(rPr))
    t = OxmlElement('w:t')
    t.text = new_text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    para.append(r)


def _make_para_from_template(template_para, text):
    """
    Deep-copy template_para, strip all runs, add a single new run with text,
    preserving the template run's rPr (font, color, size, bold …).
    """
    if template_para is None:
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p.append(r)
        return p

    p = copy.deepcopy(template_para)

    # Remove everything except pPr
    for child in list(p):
        if child.tag != f'{{{W}}}pPr':
            p.remove(child)

    if text:
        # Retrieve rPr from first run of the original
        rPr = None
        t_runs = template_para.findall(f'{{{W}}}r')
        if t_runs:
            rPr = t_runs[0].find(f'{{{W}}}rPr')

        r = OxmlElement('w:r')
        if rPr is not None:
            r.append(copy.deepcopy(rPr))
        t_el = OxmlElement('w:t')
        t_el.text = text
        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t_el)
        p.append(r)

    return p


# ── shape updaters ────────────────────────────────────────────────────────────

def _update_header_shape(doc, data):
    """Update header text box: name + up to 2 titles."""
    for _anchor, tc in _find_shapes(doc):
        text = _all_text(tc)
        if 'Maxime' in text or 'Gerard' in text or 'GERARD' in text:
            paras = tc.findall(f'{{{W}}}p')
            first  = data.get('firstName', '')
            last   = data.get('lastName', '').upper()
            titles = data.get('titles', ['', ''])

            if len(paras) >= 1:
                _replace_para_text(paras[0], f'{first} {last}'.strip())
            if len(paras) >= 2:
                _replace_para_text(paras[1], titles[0] if len(titles) > 0 else '')
            if len(paras) >= 3:
                _replace_para_text(paras[2], titles[1] if len(titles) > 1 else '')
            return


def _update_sidebar_shape(doc, data):
    """Rebuild the sidebar text box content (skills + languages)."""
    for _anchor, tc in _find_shapes(doc):
        text = _all_text(tc)
        if 'Personal skills' in text or 'Areas of expertise' in text:
            paras = tc.findall(f'{{{W}}}p')

            # Find formatting templates: one header para, one bullet para
            header_tmpl = None
            bullet_tmpl = None
            for p in paras:
                pStyle = p.find(f'.//{{{W}}}pStyle')
                style_val = (pStyle.get(f'{{{W}}}val') if pStyle is not None else None) or ''
                p_text = _all_text(p).strip()
                if not p_text:
                    continue
                if 'Paragraphe' in style_val or 'List' in style_val:
                    if bullet_tmpl is None:
                        bullet_tmpl = p
                else:
                    if header_tmpl is None:
                        header_tmpl = p
                if header_tmpl is not None and bullet_tmpl is not None:
                    break

            # Clear all paragraphs
            for p in list(tc.findall(f'{{{W}}}p')):
                tc.remove(p)

            # Build sections
            sections = []
            if data.get('personalSkills'):
                sections.append(('Personal skills', [', '.join(data['personalSkills'])]))
            if data.get('areasOfExpertise'):
                sections.append(('Areas of expertise', [', '.join(data['areasOfExpertise'])]))
            if data.get('technicalSkills'):
                sections.append(('Technical skills', data['technicalSkills']))
            if data.get('languages'):
                lang_items = [f"{l['language']}: {l['proficiency']}" for l in data['languages']]
                sections.append(('Languages', lang_items))

            for idx, (title, items) in enumerate(sections):
                if idx > 0:
                    tc.append(_make_para_from_template(header_tmpl, ''))
                tc.append(_make_para_from_template(header_tmpl, title))
                for item in items:
                    tc.append(_make_para_from_template(bullet_tmpl if bullet_tmpl is not None else header_tmpl, item))

            # Ensure at least one paragraph remains (Word requires it)
            if not tc.findall(f'{{{W}}}p'):
                tc.append(OxmlElement('w:p'))
            return


# ── body rebuild ──────────────────────────────────────────────────────────────

def _rebuild_body_content(doc, data):
    body = doc.element.body
    children = list(body)

    # Identify paragraphs that anchor floating shapes
    anchor_indices = {
        i for i, child in enumerate(children)
        if child.tag == f'{{{W}}}p' and child.findall(f'.//{{{WP}}}anchor')
    }

    # Remove everything that is not an anchor para, a sectPr, or a drawing para
    to_remove = [
        child for i, child in enumerate(children)
        if child.tag in (f'{{{W}}}p', f'{{{W}}}tbl') and i not in anchor_indices
    ]
    for child in to_remove:
        body.remove(child)

    _add_body_sections(doc, data)


def _safe_add(doc, text, style_name):
    try:
        return doc.add_paragraph(text, style=style_name)
    except Exception:
        return doc.add_paragraph(text)


def _add_body_sections(doc, data):
    # ── Professional Summary ──────────────────────────────────────────────────
    _safe_add(doc, 'Professional Summary', 'Heading 1')

    bio = data.get('bio', '')
    if bio:
        _safe_add(doc, bio, 'Normal')

    overall_dates = data.get('overallDates', '')
    if overall_dates:
        _safe_add(doc, overall_dates, 'Job Dates')

    current_role    = data.get('currentRole', '')
    current_company = data.get('currentCompany', '')
    if current_role:
        _safe_add(doc, current_role, 'Normal')
    if current_company:
        _safe_add(doc, f'At {current_company}', 'Normal')

    # ── Training & Certifications ─────────────────────────────────────────────
    certs = data.get('certifications', [])
    if certs:
        _safe_add(doc, 'Training & Certifications', 'Heading 1')
        for cert in certs:
            year = cert.get('year', '')
            name = cert.get('name', '')
            label = f'{year}: {name}' if year else name
            _safe_add(doc, label, 'List Paragraph')

    # ── Education ────────────────────────────────────────────────────────────
    education = data.get('education', [])
    if education:
        _safe_add(doc, 'Education', 'Heading 1')
        for edu in education:
            years       = edu.get('years', '')
            degree      = edu.get('degree', '')
            institution = edu.get('institution', '')
            label = f'{years}: {degree}'
            if institution:
                label += f'\n{institution}'
            _safe_add(doc, label, 'List Paragraph')

    # ── Professional Experiences ──────────────────────────────────────────────
    experiences = data.get('experiences', [])
    if experiences:
        _safe_add(doc, 'Professional Experiences', 'Heading 1')
        for exp in experiences:
            dates       = exp.get('dates', '')
            title       = exp.get('title', '')
            company     = exp.get('company', '')
            description = exp.get('description', [])
            tools       = exp.get('tools', '')

            if dates:
                _safe_add(doc, dates, 'Job Dates')

            job_line = title
            if company:
                job_line += f' – {company}'
            if job_line:
                _safe_add(doc, job_line, 'Job Details')

            for desc in description:
                if desc.strip():
                    _safe_add(doc, desc, 'List Paragraph')

            if tools:
                _safe_add(doc, f'Tools: {tools}', 'List Paragraph')


# ── public entry point ────────────────────────────────────────────────────────

def generate_beyond_data(data, template_path):
    doc = Document(template_path)
    _update_header_shape(doc, data)
    _update_sidebar_shape(doc, data)
    _rebuild_body_content(doc, data)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
